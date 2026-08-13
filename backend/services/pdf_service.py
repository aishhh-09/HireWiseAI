import os
from pypdf import PdfReader
import docx
from fpdf import FPDF

class PDFService:
    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> str:
        """Extract text content from a PDF file."""
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(docx_path: str) -> str:
        """Extract text content from a DOCX file."""
        try:
            doc = docx.Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n".join(full_text).strip()
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_txt(txt_path: str) -> str:
        """Extract text content from a TXT file."""
        try:
            with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception as e:
            raise RuntimeError(f"Failed to read TXT file: {str(e)}")

    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """Extract text from any supported format based on extension."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return cls.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return cls.extract_text_from_docx(file_path)
        elif ext in [".txt", ".md"]:
            return cls.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def convert_text_to_pdf(text: str, output_path: str, title: str = "Resume"):
        """Generate a clean, structured PDF from plain text using fpdf2."""
        try:
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                
            pdf = FPDF()
            pdf.add_page()
            
            # Use standard Helvetica font
            pdf.set_font("helvetica", "B", 16)
            pdf.cell(pdf.epw, 10, title, ln=True, align="C")
            pdf.ln(5)
            
            pdf.set_font("helvetica", size=10)
            
            # Replace tabs with spaces and encode to latin-1 to avoid fpdf2 errors
            cleaned_tabs = text.replace("\t", "    ")
            clean_text = cleaned_tabs.encode("latin-1", "replace").decode("latin-1")
            
            # Print lines
            for line in clean_text.split("\n"):
                if line.strip():
                    pdf.multi_cell(pdf.epw, 5, line)
                else:
                    pdf.ln(3)
                    
            pdf.output(output_path)
        except Exception as e:
            raise RuntimeError(f"Failed to convert text to PDF: {str(e)}")

    @staticmethod
    def compile_latex_to_pdf(parsed_data: dict, output_path: str):
        """Generates a high-quality LaTeX-style formatted PDF from parsed profile data."""
        try:
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                
            pdf = FPDF()
            pdf.add_page()
            
            # Helper to encode safe latin-1 characters
            def safe_text(txt):
                if not txt:
                    return ""
                return str(txt).replace("\t", "    ").encode("latin-1", "replace").decode("latin-1")
                
            # Set margins to 0.75 in (about 18mm)
            pdf.set_margins(18, 18, 18)
            
            personal = parsed_data.get("personal_info", {})
            name = safe_text(personal.get("name", "Candidate Name"))
            email = safe_text(personal.get("email", ""))
            phone = safe_text(personal.get("phone", ""))
            loc = safe_text(personal.get("location", ""))
            
            # 1. Header Name
            pdf.set_font("times", "B", 18)
            pdf.cell(pdf.epw, 8, name, ln=True, align="C")
            
            # 2. Contact Details Info
            contact_info = []
            if email: contact_info.append(email)
            if phone: contact_info.append(phone)
            if loc: contact_info.append(loc)
            
            pdf.set_font("times", size=10)
            contact_str = "  |  ".join(contact_info)
            pdf.cell(pdf.epw, 6, safe_text(contact_str), ln=True, align="C")
            pdf.ln(4)
            
            # Helper to draw a section header
            def add_section_header(title):
                pdf.ln(3)
                pdf.set_font("times", "B", 11)
                pdf.cell(pdf.epw, 6, safe_text(title), ln=True)
                # Draw a thin horizontal line
                pdf.set_line_width(0.2)
                pdf.line(pdf.x, pdf.y, pdf.x + pdf.epw, pdf.y)
                pdf.ln(2)
                
            # 3. Skills Section
            skills = parsed_data.get("skills", [])
            if skills:
                add_section_header("TECHNICAL SKILLS")
                pdf.set_font("times", size=10)
                skills_str = ", ".join(skills)
                pdf.multi_cell(pdf.epw, 5, safe_text(skills_str))
                pdf.ln(2)
                
            # 4. Experience Section
            experience = parsed_data.get("experience", [])
            if experience:
                add_section_header("PROFESSIONAL EXPERIENCE")
                for exp in experience:
                    pdf.set_font("times", "B", 10)
                    role = safe_text(exp.get("role", "Role"))
                    duration = safe_text(exp.get("duration", ""))
                    
                    # Print role left and duration right
                    pdf.cell(pdf.epw - 60, 5, role)
                    pdf.cell(60, 5, duration, align="R", ln=True)
                    
                    # Company in italic
                    pdf.set_font("times", "I", 10)
                    company = safe_text(exp.get("company", ""))
                    pdf.cell(pdf.epw, 5, company, ln=True)
                    
                    # Description in bullet points
                    pdf.set_font("times", size=10)
                    desc = exp.get("description", "")
                    if desc:
                        # Split by period and print as bullet points
                        sentences = [s.strip() for s in desc.split(".") if s.strip()]
                        for s in sentences:
                            bullet = "-  "
                            pdf.multi_cell(pdf.epw, 4, safe_text(f"{bullet}{s}."))
                    pdf.ln(2)
                    
            # 5. Education Section
            education = parsed_data.get("education", [])
            if education:
                add_section_header("EDUCATION")
                for edu in education:
                    pdf.set_font("times", "B", 10)
                    degree = safe_text(edu.get("degree", ""))
                    year = safe_text(edu.get("year", ""))
                    
                    pdf.cell(pdf.epw - 40, 5, degree)
                    pdf.cell(40, 5, year, align="R", ln=True)
                    
                    pdf.set_font("times", "I", 10)
                    inst = safe_text(edu.get("institution", ""))
                    pdf.cell(pdf.epw, 5, inst, ln=True)
                    pdf.ln(2)
                    
            # 6. Certifications Section
            certs = parsed_data.get("certifications", [])
            if certs:
                add_section_header("CERTIFICATIONS")
                pdf.set_font("times", size=10)
                for cert in certs:
                    pdf.multi_cell(pdf.epw, 4, safe_text(f"-  {cert}"))
                pdf.ln(2)
                
            # 7. Projects Section
            projects = parsed_data.get("projects", [])
            if projects:
                add_section_header("PROJECTS")
                for proj in projects:
                    pdf.set_font("times", "B", 10)
                    p_name = safe_text(proj.get("name", "Project"))
                    pdf.cell(pdf.epw, 5, p_name, ln=True)
                    
                    tech = proj.get("tech_stack", [])
                    if tech:
                        pdf.set_font("times", "I", 10)
                        pdf.cell(pdf.epw, 5, safe_text(f"Technologies: {', '.join(tech)}"), ln=True)
                        
                    pdf.set_font("times", size=10)
                    p_desc = proj.get("description", "")
                    if p_desc:
                        sentences = [s.strip() for s in p_desc.split(".") if s.strip()]
                        for s in sentences:
                            pdf.multi_cell(pdf.epw, 4, safe_text(f"-  {s}."))
                    pdf.ln(2)
                    
            pdf.output(output_path)
        except Exception as e:
            raise RuntimeError(f"Failed to compile LaTeX PDF: {str(e)}")



    @classmethod
    def convert_to_pdf(cls, source_path: str, output_path: str):
        """Converts docx/txt files into PDF format. If it's already a PDF, copy it or save directly."""
        ext = os.path.splitext(source_path)[1].lower()
        if ext == ".pdf":
            # If already a PDF, caller handles direct saving/moving, but let's implement simple check
            if source_path != output_path:
                import shutil
                shutil.copy(source_path, output_path)
        elif ext == ".docx":
            text = cls.extract_text_from_docx(source_path)
            cls.convert_text_to_pdf(text, output_path, title="Resume (Converted from Word)")
        elif ext in [".txt", ".md"]:
            text = cls.extract_text_from_txt(source_path)
            cls.convert_text_to_pdf(text, output_path, title="Resume (Converted from Text)")
        else:
            raise ValueError(f"Cannot convert unsupported file format: {ext}")
